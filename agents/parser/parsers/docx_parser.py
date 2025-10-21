"""
DOCX parsing using python-docx
"""

import logging
from typing import Optional
from docx import Document

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse DOCX CVs using python-docx"""

    def __init__(self):
        self.logger = logger

    def parse(self, file_path: str) -> Optional[str]:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text or None if parsing fails
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Preserve some formatting info
                    if para.style.name.startswith('Heading'):
                        text_parts.append(f"\n\n## {text} ##\n")
                    else:
                        text_parts.append(text)
                        text_parts.append("\n")

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\n[Table {table_idx + 1}]\n")

                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                        text_parts.append("\n")

                text_parts.append("\n")

            result = "".join(text_parts)

            if not result or len(result.strip()) < 50:
                self.logger.warning("Extracted text too short, might be empty document")
                return None

            return result

        except Exception as e:
            self.logger.error(f"DOCX parsing failed: {e}")
            return None

    def get_metadata(self, file_path: str) -> dict:
        """Extract DOCX metadata"""
        metadata = {}

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            metadata = {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
            }

        except Exception as e:
            self.logger.warning(f"Could not extract metadata: {e}")

        return metadata
